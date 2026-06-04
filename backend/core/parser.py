"""Document parsers: txt / md / docx / pdf."""
from __future__ import annotations

import io
from pathlib import Path

from pypdf import PdfReader
from docx import Document


SUPPORTED_EXTS = {".txt", ".md", ".markdown", ".pdf", ".docx"}


def parse_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix == ".docx":
        return _parse_docx(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def _parse_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = _normalize_pdf_text(text)
        if text.strip():
            pages.append(f"[Page {i+1}]\n{text}")
    return "\n\n".join(pages)


def _parse_docx(path: Path) -> str:
    doc = Document(str(path))
    blocks: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text)
    # 表格可能存在格式异常（如缺少 <w:tblGrid>），捕获异常避免整篇文档被丢弃
    try:
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
    except Exception:
        pass
    return "\n".join(blocks)


def _normalize_pdf_text(text: str) -> str:
    lines = [ln.strip() for ln in text.splitlines()]
    lines = [ln for ln in lines if ln]
    merged: list[str] = []
    buf = ""
    for ln in lines:
        if buf and not buf.endswith(("。", "！", "？", ".", "!", "?", ":", "：")):
            buf += ln
        else:
            if buf:
                merged.append(buf)
            buf = ln
    if buf:
        merged.append(buf)
    return "\n".join(merged)
